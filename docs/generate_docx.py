"""Generate the QA Service technical documentation as a Word (.docx) file.

Run from the assignment root::

    python docs/generate_docx.py

Produces ``docs/QA-Service-Documentation.docx``.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# Palette
INK = RGBColor(0x1F, 0x2A, 0x37)
ACCENT = RGBColor(0x2B, 0x6C, 0xB0)
MUTED = RGBColor(0x55, 0x5F, 0x6B)
CODE_BG = "F2F4F7"
HDR_BG = "2B6CB0"


def _shade(cell, fill: str) -> None:
    """Apply a solid background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _code_block(doc: Document, code: str) -> None:
    """Render a monospaced, shaded code block."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _shade(cell, CODE_BG)
    cell.paragraphs[0].text = ""
    for i, line in enumerate(code.strip("\n").split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = INK
    doc.add_paragraph()


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Render a styled table with a coloured header row."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        _shade(hdr[i], HDR_BG)
        para = hdr[i].paragraphs[0]
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r, row in enumerate(rows):
        cells = table.add_row().cells
        for i, text in enumerate(row):
            if r % 2 == 1:
                _shade(cells[i], "EEF3F9")
            para = cells[i].paragraphs[0]
            run = para.add_run(text)
            run.font.size = Pt(9.5)
            run.font.color.rgb = INK
    doc.add_paragraph()


def _h(doc: Document, text: str, level: int) -> None:
    """Add a heading with consistent colour styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = ACCENT if level <= 1 else INK


def _body(doc: Document, text: str) -> None:
    """Add a normal body paragraph."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build() -> Path:
    """Build the document and return the output path."""
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK

    # ---- Cover ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("QA Service")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("RAG Document Q&A — Technical Documentation")
    run.font.size = Pt(15)
    run.font.color.rgb = MUTED

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Version 0.1.0  ·  FastAPI · ChromaDB · NVIDIA NIM · sentence-transformers")
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED
    doc.add_paragraph()

    # ---- 1. Overview ----
    _h(doc, "1. Overview", 1)
    _body(
        doc,
        "QA Service is a containerized FastAPI application that ingests PDF documents, answers "
        "natural-language questions with citations using Retrieval-Augmented Generation (RAG), and "
        "evaluates answer quality with an LLM-as-judge. Each ingested PDF is stored in its own "
        "ChromaDB collection, so a single document behaves as an independently queryable vector "
        "database. Queries may target a specific document via doc_id or fan out across every indexed "
        "document.",
    )
    _h(doc, "Key Capabilities", 2)
    _bullets(
        doc,
        [
            "PDF ingestion with token-aware recursive chunking and local embeddings.",
            "Grounded question answering with source citations (RAG).",
            "Automated answer-quality scoring via an LLM-as-judge (/evaluate).",
            "Provider-agnostic LLM access through any OpenAI-compatible endpoint.",
            "Persistent on-disk vector storage (no external database container required).",
            "Structured JSON logging (structlog) and Prometheus metrics.",
        ],
    )

    # ---- 2. Architecture ----
    _h(doc, "2. Architecture", 1)
    _body(
        doc,
        "The service is a single FastAPI process. Embeddings are computed locally on CPU with "
        "sentence-transformers; the chat model is reached over an OpenAI-compatible HTTP API "
        "(NVIDIA NIM by default); vectors are persisted to a local ChromaDB directory.",
    )
    _code_block(
        doc,
        """
+--------------------------------------------------------------+
|  qa-service  (FastAPI :8000)                                 |
|   /ingest   /query   /documents   /evaluate   /metrics       |
|        |                                                     |
|        +-- sentence-transformers  (local CPU embeddings)     |
|        +-- NVIDIA NIM             (LLM, OpenAI-compatible)   |
|        +-- ChromaDB               (persistent on-disk)       |
+--------------------------------------------------------------+
""",
    )
    _h(doc, "Per-Document Collections", 2)
    _body(
        doc,
        "Each document is written to its own Chroma collection named doc_<doc_id> (hyphens removed). "
        "Collection metadata records the doc_id, title, ingest timestamp, and chunk count. Scoping a "
        "query to a doc_id reads a single collection; omitting it fans the query out across all "
        "collections and merges the top-ranked chunks.",
    )

    # ---- 3. Technology Stack ----
    _h(doc, "3. Technology Stack", 1)
    _table(
        doc,
        ["Component", "Technology", "Role"],
        [
            ["API framework", "FastAPI + Uvicorn", "HTTP routing, OpenAPI docs, async I/O"],
            ["Validation", "Pydantic v2 / pydantic-settings", "Request/response models and config"],
            ["Embeddings", "sentence-transformers (all-MiniLM-L6-v2)", "Local CPU vectors (dim 384)"],
            ["LLM", "NVIDIA NIM via openai SDK", "Chat completions (OpenAI-compatible)"],
            ["Vector store", "ChromaDB (PersistentClient/HttpClient)", "Per-document vector collections"],
            ["Chunking", "tiktoken (cl100k_base)", "Token-aware recursive splitting"],
            ["PDF parsing", "pypdf", "Text extraction"],
            ["Logging", "structlog", "Structured JSON logs"],
            ["Metrics", "prometheus-client", "Counters exposed at /metrics"],
            ["Runtime", "Python 3.12+, Docker", "Multi-stage slim image, non-root user"],
        ],
    )

    # ---- 4. API Reference ----
    _h(doc, "4. API Reference", 1)
    _body(doc, "The service listens on port 8000. Interactive OpenAPI docs are served at /docs.")
    _table(
        doc,
        ["Method", "Path", "Description"],
        [
            ["POST", "/ingest", "Multipart PDF upload, indexed into Chroma"],
            ["POST", "/query", "Ask a question, get an answer plus citations"],
            ["GET", "/documents", "List indexed documents"],
            ["POST", "/evaluate", "Score generated answers against expected answers"],
            ["GET", "/metrics", "Prometheus exposition of service counters"],
            ["GET", "/health", "Liveness probe"],
        ],
    )

    _h(doc, "4.1  POST /ingest", 2)
    _body(
        doc,
        "Accepts a multipart upload (file plus optional title). The endpoint validates the "
        "content type / extension (415 on non-PDF), rejects empty uploads (400), enforces a 25 MB "
        "size cap (413), and returns 400 if the PDF yields no extractable text. On success it "
        "returns 201 Created with the new doc_id and chunk count.",
    )
    _code_block(
        doc,
        """
curl -X POST http://localhost:8000/ingest \\
  -F "file=@/path/to/handbook.pdf" \\
  -F "title=Company Handbook"

# 201 Created
{ "doc_id": "5b3a...-uuid",
  "title": "Company Handbook",
  "chunks_indexed": 12,
  "status": "success" }
""",
    )

    _h(doc, "4.2  POST /query", 2)
    _body(
        doc,
        "Embeds the question, retrieves the top_k most similar chunks (default 4), and asks the LLM "
        "to answer using only the retrieved context. When no relevant chunks are found the service "
        "returns a safe 'I don't know based on the provided documents.' answer. Pass an optional "
        "doc_id to scope the query to a single document.",
    )
    _code_block(
        doc,
        """
curl -X POST http://localhost:8000/query \\
  -H "Content-Type: application/json" \\
  -d '{"question":"What is the leave policy?"}'

{ "answer": "Employees are entitled to 18 days of paid leave...",
  "sources": [ {"doc_id": "uuid", "title": "Company Handbook", "chunk": "..."} ] }
""",
    )

    _h(doc, "4.3  GET /documents", 2)
    _body(doc, "Returns a summary of every indexed document (doc_id, title, chunk count, ingest time).")

    _h(doc, "4.4  POST /evaluate", 2)
    _body(
        doc,
        "Runs each test case through the RAG pipeline and scores the generated answer against the "
        "expected answer using the LLM-as-judge. The judge returns a 0.0-1.0 score plus reasoning, "
        "and the response aggregates per-case results with an average score. The request is bounded "
        "to a maximum of 50 test cases (HTTP 422 above the cap) to prevent unbounded LLM cost / "
        "denial-of-service amplification, since each case triggers two upstream LLM calls.",
    )
    _code_block(
        doc,
        """
curl -X POST http://localhost:8000/evaluate \\
  -H "Content-Type: application/json" \\
  -d '{ "test_cases": [
          {"question":"What is the notice period?","expected_answer":"30 days"}
       ] }'
""",
    )

    _h(doc, "4.5  GET /metrics", 2)
    _body(
        doc,
        "Prometheus exposition format exposing documents_ingested_total and queries_processed_total. "
        "The endpoint is documented in the OpenAPI schema under the 'observability' tag.",
    )

    # ---- 5. RAG Pipeline ----
    _h(doc, "5. RAG Pipeline", 1)
    _h(doc, "5.1  Ingestion Flow", 2)
    _bullets(
        doc,
        [
            "Extract text from the PDF page-by-page (pypdf); skip pages with no extractable text.",
            "Chunk the text with a token-aware recursive splitter (tiktoken cl100k_base).",
            "Embed each chunk locally with sentence-transformers (all-MiniLM-L6-v2, dim 384).",
            "Create a per-document Chroma collection and upsert chunks with metadata.",
            "Increment documents_ingested_total.",
        ],
    )
    _h(doc, "5.2  Query Flow", 2)
    _bullets(
        doc,
        [
            "Embed the incoming question.",
            "Retrieve the top_k nearest chunks (single collection if doc_id given, else fan-out).",
            "Format the retrieved chunks as numbered, titled context.",
            "Prompt the LLM (system + user templates from YAML) to answer using only that context.",
            "Return the answer with source citations; increment queries_processed_total.",
        ],
    )
    _h(doc, "5.3  Chunking Strategy", 2)
    _body(
        doc,
        "The chunker splits recursively on a separator hierarchy (paragraphs, lines, sentence "
        "punctuation, clauses, words) so chunks respect natural boundaries while staying within the "
        "token budget. Any piece still over budget is hard-split on token IDs. Consecutive chunks "
        "carry a configurable token overlap (default 50) to preserve context across boundaries.",
    )
    _h(doc, "5.4  LLM-as-Judge", 2)
    _body(
        doc,
        "The judge prompts the model for a strict JSON verdict {score, reasoning}. Responses are "
        "parsed tolerantly (the first JSON object is extracted and the score clamped to [0, 1]). On "
        "a parse failure the judge is retried once at a higher temperature; if that also fails a "
        "deterministic fallback verdict (score 0.0) is returned so the batch can complete.",
    )

    # ---- 6. Configuration ----
    _h(doc, "6. Configuration", 1)
    _body(
        doc,
        "All settings are loaded from environment variables (or a local .env file in development) "
        "via pydantic-settings. The .env is resolved by absolute path so local runs work regardless "
        "of the working directory; inside Docker the variables are injected directly.",
    )
    _table(
        doc,
        ["Variable", "Default", "Purpose"],
        [
            ["LLM_PROVIDER", "nvidia", "LLM provider selector"],
            ["LLM_BASE_URL", "https://integrate.api.nvidia.com/v1", "OpenAI-compatible base URL"],
            ["LLM_MODEL", "meta/llama-3.3-70b-instruct", "Chat model name"],
            ["LLM_API_KEY", "(required)", "Provider API key (free NVIDIA tier)"],
            ["LLM_TEMPERATURE", "0.1", "Sampling temperature (0.0-2.0)"],
            ["LLM_MAX_TOKENS", "512", "Response token cap (1-8192)"],
            ["EMBEDDING_MODEL", "all-MiniLM-L6-v2", "Local embedding model"],
            ["VECTOR_STORE_URL", "./.chroma (local)", "Chroma dir or http:// endpoint"],
            ["CHUNK_SIZE", "512", "Target tokens per chunk (64-4096)"],
            ["CHUNK_OVERLAP", "50", "Token overlap between chunks (0-512)"],
            ["TOP_K", "4", "Retrieved chunks per query (1-20)"],
            ["LOG_LEVEL", "INFO", "structlog level"],
            ["APP_PORT", "8000", "Server port"],
        ],
    )
    _h(doc, "Swapping LLM Providers", 2)
    _body(
        doc,
        "Because the LLM goes through any OpenAI-compatible endpoint, the provider is changed purely "
        "by configuration. NVIDIA NIM, OpenAI, and local Ollama (http://host.docker.internal:11434/v1, "
        "key 'not-needed') are all supported by setting LLM_BASE_URL, LLM_MODEL, and LLM_API_KEY.",
    )

    # ---- 7. Running the Service ----
    _h(doc, "7. Running the Service", 1)
    _h(doc, "7.1  Docker Compose", 2)
    _code_block(
        doc,
        """
cd assignments/qa-service
cp .env.example .env
# set LLM_API_KEY (free key at https://build.nvidia.com/)
docker compose up --build
# -> http://localhost:8000  (docs at /docs)
""",
    )
    _body(
        doc,
        "Chroma is embedded and persistent: the compose file runs a single qa-service container with "
        "VECTOR_STORE_URL=/app/.chroma backed by a named volume, so no separate database container is "
        "required.",
    )
    _h(doc, "7.2  Local Development", 2)
    _code_block(
        doc,
        """
cd assignments/qa-service
python -m venv .venv
.venv\\Scripts\\activate          # Windows
pip install -e ".[dev]"

# point Chroma at a local directory and add your key
echo VECTOR_STORE_URL=./.chroma >> .env
echo LLM_API_KEY=<your-nvidia-key> >> .env

uvicorn qa_service.main:app --reload --port 8000
""",
    )

    # ---- 8. Project Layout ----
    _h(doc, "8. Project Layout", 1)
    _code_block(
        doc,
        """
qa-service/
  Dockerfile                 # multi-stage slim runtime, non-root user
  docker-compose.yml         # single service, persistent on-disk Chroma
  .env.example               # all env vars documented
  pyproject.toml             # dependencies + tooling config
  src/qa_service/
    main.py                  # FastAPI app + /metrics route
    config.py                # pydantic-settings
    dependencies.py          # DI singletons
    models/                  # request/response models
    routes/                  # API endpoints
    services/                # rag, judge, chunker, pdf_parser
    providers/               # embedding + LLM adapters
    repository/              # Chroma vector store
    observability/           # structlog + Prometheus
    prompts/                 # YAML prompt templates
  tests/                     # unit + integration tests
""",
    )

    # ---- 9. Observability ----
    _h(doc, "9. Observability", 1)
    _bullets(
        doc,
        [
            "Logging: structlog emits structured JSON logs (configurable level) with event names "
            "such as service_starting, ingest_complete, and evaluation_complete.",
            "Metrics: prometheus-client exposes documents_ingested_total and "
            "queries_processed_total at /metrics; noisy *_created series are disabled.",
            "Health: GET /health provides a liveness probe used by the Docker healthcheck.",
        ],
    )

    # ---- 10. Security ----
    _h(doc, "10. Security Considerations", 1)
    _bullets(
        doc,
        [
            "Upload hardening: content-type/extension validation, 25 MB size cap, empty-file and "
            "unparseable-PDF rejection.",
            "Cost / DoS control: /evaluate caps test_cases at 50 (HTTP 422 above) since each case "
            "triggers two LLM calls.",
            "Secrets: the API key is supplied via environment variables / .env (gitignored) and is "
            "never hardcoded or logged.",
            "Container: runs as a non-root user on a slim multi-stage image.",
            "Deployment notes: authentication and rate limiting are intentionally left to the "
            "deployment layer (gateway / reverse proxy) and are out of scope for this service.",
        ],
    )

    out = Path(__file__).resolve().parent / "QA-Service-Documentation.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
